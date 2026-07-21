# This module contains part of the application's Python implementation.

"""Export selected workspace files into a single Word (.docx) document."""  # Module description

"""This script collects text-based files from the project directory, reads
their contents using a tolerant encoding strategy, sanitizes nonprintable
characters, and writes each file's contents to a headinged section in
the generated Word document.

Do not change executable behavior; comments only."""  # Extended explanation and warning

from pathlib import Path  # Import Path for filesystem path manipulation

from docx import Document  # Import Document to create .docx files

root = Path(".").resolve()  # Resolve the project root to an absolute Path
out_path = root / "project_export.docx"  # Destination file for the export

include_exts = {  # Set of file extensions to include in the export
    ".py",  # Python source
    ".html",  # HTML templates
    ".css",  # Stylesheets
    ".js",  # JavaScript files
    ".ts",  # TypeScript files
    ".json",  # JSON data/config
    ".md",  # Markdown files
    ".txt",  # Plain text files
    ".yml",  # YAML files
    ".yaml",  # YAML alternative
    ".env",  # Environment files
    ".xml",  # XML files
}  # end include_exts

files = []  # Initialize list to collect matching files
for path in root.rglob("*"):  # Recursively walk all filesystem entries
    if (
        path.is_file() and path.suffix.lower() in include_exts
    ):  # File and allowed extension
        if any(
            part in {"__pycache__", ".venv", "node_modules", ".git"}
            for part in path.parts
        ):  # Skip common folders
            continue  # Skip unwanted directories
        files.append(path)  # Add matching file to list

files = sorted(files, key=lambda p: str(p).lower())  # Sort files for predictable order

doc = Document()  # Create a new Word document
doc.add_heading("Project Export", level=1)  # Add top-level heading
doc.add_paragraph("Generated from the workspace contents.")  # Add a summary paragraph

for path in files:  # Iterate over collected files
    rel = path.relative_to(
        root
    ).as_posix()  # Compute POSIX-style relative path for headings
    doc.add_heading(rel, level=2)  # Add a heading for the file path
    try:
        text = path.read_text(encoding="utf-8")  # Try reading file as UTF-8 text
    except Exception:  # If reading as UTF-8 fails
        try:
            text = path.read_text(encoding="latin-1")  # Fall back to Latin-1 encoding
        except Exception:
            text = "[Binary or unreadable content]"  # Mark as unreadable if both attempts fail

    for line in text.splitlines():  # Iterate lines in file content
        safe_line = "".join(
            ch for ch in line if ch in "\t\n\r" or ch.isprintable()
        )  # Keep printable characters
        safe_line = safe_line.replace("\x00", "")  # Remove null bytes
        if safe_line.strip():  # If the line contains visible characters
            doc.add_paragraph(
                safe_line, style="List Bullet"
            )  # Add as a bulleted paragraph
        else:
            doc.add_paragraph("")  # Preserve blank lines in output

doc.save(out_path)  # Save the assembled Word document to disk
print(out_path)  # Print the output path to stdout for convenience
